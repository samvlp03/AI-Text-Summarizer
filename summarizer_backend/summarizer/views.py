from rest_framework import status, generics
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from .models import Summary, ExportFormat
from .serializers import SummarySerializer, ExportFormatSerializer, UserSerializer, MyTokenObtainPairSerializer
from django.contrib.auth.models import User
from django.http import HttpResponse
import ollama
import json
from docx import Document
from io import BytesIO
import pdfkit
from rest_framework_simplejwt.views import TokenObtainPairView

MODEL = 'llama3.2:1b'

class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer

class UserCreateView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = UserSerializer

class SummarizeView(APIView):
    permission_classes = [IsAuthenticated]
    
    def _validate_input(self, text, focus):
        if len(text) > 8000:
            raise ValueError("Input text too long (max 8000 characters)")
        if len(focus) > 80:
            raise ValueError("Focus phrase too long (max 80 characters)")
        return text[:8000], focus[:80]
    
    def _build_prompt(self, length, tonality, focus):
        return f"""**Llama 3.2 1B Summary Instructions**

TASK: Create a {length} summary with {tonality} tone.

PARAMETERS:
- Length: {{
    'short': 'STRICTLY 1-2 SENTENCES',
    'medium': 'EXACTLY 3-5 SENTENCES', 
    'long': 'PRECISELY 6-8 SENTENCES'
}}[length]
- Style: {tonality.upper()}
- Focus: {'"' + focus + '"' if focus else 'All main points'}

RULES:
1. NEVER EXCEED SENTENCE COUNT
2. USE SIMPLE LANGUAGE
3. PRESERVE FACTS ONLY
4. OMIT EXAMPLES/ANALOGIES
5. {"PRIORITIZE: " + focus.upper() if focus else ""}"""
    
    def post(self, request):
        try:
            # Validate and prepare input
            text, focus = self._validate_input(
                request.data.get('text', ''),
                request.data.get('focus', '')
            )
            
            # Get parameters with model-appropriate defaults
            params = {
                'length': request.data.get('length', 'medium'),
                'tonality': request.data.get('tonality', 'neutral'),
                'temperature': min(float(request.data.get('temperature', 0.65)), 0.75),
                'top_p': max(min(float(request.data.get('top_p', 0.9)), 0.7), 0.95),
                'focus': focus
            }
            
            # Generate with Llama 3.2 1B
            response = ollama.chat(
                model=MODEL,
                messages=[
                    {'role': 'system', 'content': self._build_prompt(**params)},
                    {'role': 'user', 'content': f"TEXT:\n{text}"}
                ],
                options={
                    'temperature': params['temperature'],
                    'top_p': params['top_p'],
                    'num_ctx': 4096,
                    'num_predict': 256,
                    'repeat_penalty': 1.15
                }
            )
            
            # Process and validate output
            summary = self._process_response(response, params['length'])
            
            # Save result
            record = Summary.objects.create(
                user=request.user,
                original_text=text,
                summary_text=summary,
                **params,
                model_used=MODEL
            )
            
            return Response(SummarySerializer(record).data)
            
        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST if isinstance(e, ValueError) 
                       else status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _process_response(self, response, length):
        """Validate and format the model output"""
        text = response['message']['content'].strip()
        
        # Basic validation
        if not text or len(text.split()) < 4:
            raise ValueError("Insufficient summary generated")
        
        # Enforce length requirements
        sentences = [s.strip() for s in text.split('. ') if s.strip()]
        sentence_limits = {'short': 2, 'medium': 5, 'long': 8}
        
        if len(sentences) > sentence_limits.get(length, 5):
            text = '. '.join(sentences[:sentence_limits[length]]) + '.'
        
        return text
    
class RegenerateView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, summary_id):
        try:
            # Validate and prepare input
            text, focus = self._validate_input(
                request.data.get('text', ''),
                request.data.get('focus', '')
            )
            
            # Get parameters with model-appropriate defaults
            params = {
                'length': request.data.get('length', 'medium'),
                'tonality': request.data.get('tonality', 'neutral'),
                'temperature': min(float(request.data.get('temperature', 0.65)), 0.75),
                'top_p': max(min(float(request.data.get('top_p', 0.9)), 0.7), 0.95),
                'focus': focus
            }
            
            summary = Summary.objects.get(id=summary_id, user=request.user)
            
            # Same system message as SummarizeView
            system_message = f"""**SUMMARY GENERATION PROTOCOL**

MODEL: Llama 3.2 1B (Precision Mode)

INPUT ANALYSIS:
- Original text length: {len(text.split())} words
- Key focus: {'"' + focus + '"' if focus else 'All main points'}
- Output style: {params["tonality"].upper()}

GENERATION PARAMETERS:
- Length class: {params["length"].upper()}
- Temperature: {params["temperature"]}
- Top-p: {params["top_p"]}

EXECUTION RULES:
1. STRICTLY FOLLOW LENGTH REQUIREMENT
2. PRESERVE ORIGINAL MEANING
3. EXCLUDE EXAMPLES/ANALOGIES
4. USE SIMPLE SENTENCE STRUCTURES
5. {"FOCUS ON: " + focus.upper() if focus else "COVER ALL ESSENTIAL ELEMENTS"}

OUTPUT FORMAT:
- Plain text only
- No bullet points
- Complete sentences
- No meta-commentary"""
            
            response = ollama.chat(
                model=MODEL,
                messages=[
                    {'role': 'system', 'content': system_message},
                    {'role': 'user', 'content': f"Summarize this text:\n{summary.original_text}"}
                ],
                options={
                    'temperature': float(request.data.get('temperature', summary.temperature)),
                    'top_p': float(request.data.get('top_p', summary.top_p)),
                    'num_ctx': 4096
                }
            )
            
            # Rest of the regeneration logic remains the same
            summary.summary_text = response['message']['content']
            summary.length = params["length"]
            summary.tonality = params["tonality"]
            summary.temperature = params["temperature"]
            summary.top_p = params["top_p"]
            summary.focus = focus
            summary.save()
            
            serializer = SummarySerializer(summary)
            return Response(serializer.data, status=status.HTTP_200_OK)
        except Summary.DoesNotExist:
            return Response({'error': 'Summary not found!'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class UpdateSummaryView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request, summary_id):
        try:
            summary = Summary.objects.get(id=summary_id, user=request.user)
            summary.summary_text = request.data.get('summary_text', summary.summary_text)
            summary.save()
            
            serializer = SummarySerializer(summary)
            return Response(serializer.data, status=status.HTTP_200_OK)
            
        except Summary.DoesNotExist:
            return Response({'error': 'Summary not found!!'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ToggleFavoriteView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, summary_id):
        try:
            summary = Summary.objects.get(id=summary_id, user=request.user)
            summary.is_favorite = not summary.is_favorite
            summary.save()
            
            serializer = SummarySerializer(summary)
            return Response(serializer.data, status=status.HTTP_200_OK)
            
        except Summary.DoesNotExist:
            return Response({'error': 'Summary not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class SummaryListView(generics.ListAPIView):
    serializer_class = SummarySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = Summary.objects.filter(user=user).order_by('-created_at')
        
        is_favorite = self.request.query_params.get('is_favorite', None)
        if is_favorite is not None:
            queryset = queryset.filter(is_favorite=is_favorite.lower() == 'true')
            
        return queryset

class ExportSummaryView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, summary_id):
        try:
            summary = Summary.objects.get(id=summary_id, user=request.user)
            format_type = request.query_params.get('format', 'pdf')
            
            if format_type == 'pdf':
                html = f"""
                <h1>Summary</h1>
                <h3>Original Text</h3>
                <p>{summary.original_text}</p>
                <h3>Summary</h3>
                <p>{summary.summary_text}</p>
                <p>Generated on {summary.created_at}</p>
                """
                pdf = pdfkit.from_string(html, False)
                response = HttpResponse(pdf, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="summary_{summary_id}.pdf"'
                return response
                
            elif format_type == 'docx':
                document = Document()
                document.add_heading('Summary', 0)
                document.add_heading('Original Text', level=1)
                document.add_paragraph(summary.original_text)
                document.add_heading('Summary', level=1)
                document.add_paragraph(summary.summary_text)
                document.add_paragraph(f"Generated on {summary.created_at}")
                
                file_stream = BytesIO()
                document.save(file_stream)
                file_stream.seek(0)
                
                response = HttpResponse(
                    file_stream.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="summary_{summary_id}.docx"'
                return response
                
            elif format_type == 'json':
                data = {
                    'original_text': summary.original_text,
                    'summary_text': summary.summary_text,
                    'created_at': summary.created_at
                }
                response = HttpResponse(json.dumps(data), content_type='application/json')
                response['Content-Disposition'] = f'attachment; filename="summary_{summary_id}.json"'
                return response
                
            else:
                return Response({'error': 'Unsupported format'}, status=status.HTTP_400_BAD_REQUEST)
                
        except Summary.DoesNotExist:
            return Response({'error': 'Summary not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ExportFormatListView(generics.ListAPIView):
    queryset = ExportFormat.objects.all()
    serializer_class = ExportFormatSerializer
    permission_classes = [IsAuthenticated]